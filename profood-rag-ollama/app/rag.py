from __future__ import annotations

import csv
import json
import shutil
import time
from collections.abc import Iterator
from pathlib import Path
from threading import Event
from typing import Any

from docx import Document as WordDocument
from langchain_chroma import Chroma
from langchain_community.document_loaders import PyPDFDirectoryLoader
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_ollama import ChatOllama, OllamaEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from openpyxl import load_workbook

from app.config import settings


PROMPT = ChatPromptTemplate.from_messages(
    [
        (
            "system",
            "You are Profood AI, a helpful assistant for a Moroccan marketplace app "
            "that contains food products, agriculture products, equipment, suppliers, "
            "articles, and forum knowledge.\n\n"
            "Rules:\n"
            "1. Answer using ONLY the context provided.\n"
            "2. If the context is not enough, say that Profood does not have enough data yet.\n"
            "3. Be practical and clear.\n"
            "4. When useful, mention which product/equipment/forum source supports the answer.\n"
            "5. Do not invent prices, sellers, or availability.\n\n"
            "Context:\n{context}",
        ),
        ("human", "Question: {question}"),
    ]
)

VOICE_PROMPT = ChatPromptTemplate.from_messages(
    [
        (
            "system",
            "You are Profood AI, a helpful assistant for a Moroccan marketplace app "
            "that contains food products, agriculture products, equipment, suppliers, "
            "articles, and forum knowledge.\n\n"
            "Rules:\n"
            "1. Answer using ONLY the context provided.\n"
            "2. If the context is not enough, say that Profood does not have enough data yet.\n"
            "3. Be practical and clear.\n"
            "4. When useful, mention which product/equipment/forum source supports the answer.\n"
            "5. Do not invent prices, sellers, or availability.\n"
            "6. Answer in French. Use 2 to 4 short sentences. Be direct, clear, and suitable for spoken audio.\n\n"
            "Context:\n{context}",
        ),
        ("human", "Question: {question}"),
    ]
)

_EMBEDDINGS: OllamaEmbeddings | None = None
_LLM: ChatOllama | None = None
_VECTOR_STORE: Chroma | None = None


def _project_root() -> Path:
    return Path(__file__).resolve().parents[1]


def _resolve_path(path_value: str) -> Path:
    path = Path(path_value)

    if path.is_absolute():
        return path

    return _project_root() / path


def get_embeddings() -> OllamaEmbeddings:
    global _EMBEDDINGS

    if _EMBEDDINGS is None:
        _EMBEDDINGS = OllamaEmbeddings(
            model=settings.ollama_embedding_model,
            base_url=settings.ollama_base_url,
        )

    return _EMBEDDINGS


def get_llm() -> ChatOllama:
    global _LLM

    if _LLM is None:
        _LLM = ChatOllama(
            model=settings.ollama_chat_model,
            base_url=settings.ollama_base_url,
            temperature=0.1,
        )

    return _LLM


def get_vector_store() -> Chroma:
    global _VECTOR_STORE

    chroma_dir = _resolve_path(settings.chroma_dir)
    chroma_dir.mkdir(parents=True, exist_ok=True)

    if _VECTOR_STORE is None:
        _VECTOR_STORE = Chroma(
            collection_name=settings.collection_name,
            embedding_function=get_embeddings(),
            persist_directory=str(chroma_dir),
        )

    return _VECTOR_STORE


def _guess_doc_type(source_path: str) -> str:
    name = Path(source_path).name.lower()

    if "equipment" in name or "equipement" in name:
        return "equipment"

    if "forum" in name or "faq" in name:
        return "forum"

    if "product" in name or "produit" in name:
        return "product"

    if "supplier" in name or "fournisseur" in name:
        return "supplier"

    if "article" in name or "guide" in name:
        return "article"

    return "document"


def _dict_to_text(data: dict[str, Any]) -> str:
    """
    Convert a dictionary row/object into readable text for RAG.

    Example:
    {
        "name": "Olive oil press",
        "type": "equipment"
    }

    becomes:

    name: Olive oil press
    type: equipment
    """

    lines: list[str] = []

    for key, value in data.items():
        if value is None:
            continue

        if isinstance(value, (dict, list)):
            value = json.dumps(value, ensure_ascii=False)

        value_as_text = str(value).strip()

        if value_as_text:
            lines.append(f"{key}: {value_as_text}")

    return "\n".join(lines)


def _load_text_file(file_path: Path) -> list[Document]:
    """
    Load simple text files:
    - .txt
    - .md

    One file becomes one LangChain Document.
    """

    content = file_path.read_text(encoding="utf-8").strip()

    if not content:
        return []

    return [
        Document(
            page_content=content,
            metadata={
                "source": str(file_path),
                "source_file": file_path.name,
                "doc_type": _guess_doc_type(file_path.name),
                "file_type": file_path.suffix.lower(),
            },
        )
    ]


def _load_json_file(file_path: Path) -> list[Document]:
    """
    Load JSON files.

    Supported shapes:

    1. List of objects:
    [
        {"type": "product", "name": "Olive oil"},
        {"type": "equipment", "name": "Olive oil press"}
    ]

    2. Single object:
    {
        "type": "article",
        "title": "Olive oil guide"
    }
    """

    content = file_path.read_text(encoding="utf-8")
    data = json.loads(content)

    docs: list[Document] = []

    if isinstance(data, list):
        for index, item in enumerate(data, start=1):
            if isinstance(item, dict):
                page_content = _dict_to_text(item)
                doc_type = str(item.get("type") or item.get("doc_type") or _guess_doc_type(file_path.name))
            else:
                page_content = str(item)
                doc_type = _guess_doc_type(file_path.name)

            if not page_content.strip():
                continue

            docs.append(
                Document(
                    page_content=page_content,
                    metadata={
                        "source": str(file_path),
                        "source_file": file_path.name,
                        "doc_type": doc_type,
                        "file_type": ".json",
                        "record_index": index,
                    },
                )
            )

    elif isinstance(data, dict):
        page_content = _dict_to_text(data)
        doc_type = str(data.get("type") or data.get("doc_type") or _guess_doc_type(file_path.name))

        if page_content.strip():
            docs.append(
                Document(
                    page_content=page_content,
                    metadata={
                        "source": str(file_path),
                        "source_file": file_path.name,
                        "doc_type": doc_type,
                        "file_type": ".json",
                    },
                )
            )

    else:
        page_content = str(data).strip()

        if page_content:
            docs.append(
                Document(
                    page_content=page_content,
                    metadata={
                        "source": str(file_path),
                        "source_file": file_path.name,
                        "doc_type": _guess_doc_type(file_path.name),
                        "file_type": ".json",
                    },
                )
            )

    return docs


def _load_csv_file(file_path: Path) -> list[Document]:
    """
    Load CSV files.

    Each row becomes one LangChain Document.
    """

    docs: list[Document] = []

    with file_path.open("r", encoding="utf-8-sig", newline="") as csv_file:
        reader = csv.DictReader(csv_file)

        for index, row in enumerate(reader, start=1):
            page_content = _dict_to_text(row)

            if not page_content.strip():
                continue

            doc_type = str(row.get("type") or row.get("doc_type") or _guess_doc_type(file_path.name))

            docs.append(
                Document(
                    page_content=page_content,
                    metadata={
                        "source": str(file_path),
                        "source_file": file_path.name,
                        "doc_type": doc_type,
                        "file_type": ".csv",
                        "record_index": index,
                    },
                )
            )

    return docs


def _load_docx_file(file_path: Path) -> list[Document]:
    """
    Load Word .docx files.

    It reads:
    - normal paragraphs
    - tables
    """

    word_doc = WordDocument(str(file_path))

    parts: list[str] = []

    for paragraph in word_doc.paragraphs:
        text = paragraph.text.strip()

        if text:
            parts.append(text)

    for table_index, table in enumerate(word_doc.tables, start=1):
        rows_text: list[str] = []

        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(cells).strip()

            if row_text:
                rows_text.append(row_text)

        if rows_text:
            parts.append(f"Table {table_index}:\n" + "\n".join(rows_text))

    content = "\n\n".join(parts).strip()

    if not content:
        return []

    return [
        Document(
            page_content=content,
            metadata={
                "source": str(file_path),
                "source_file": file_path.name,
                "doc_type": _guess_doc_type(file_path.name),
                "file_type": ".docx",
            },
        )
    ]


def _load_excel_file(file_path: Path) -> list[Document]:
    """
    Load Excel files.

    Supported:
    - .xlsx
    - .xlsm
    - .xltx
    - .xltm

    Each sheet becomes one LangChain Document.
    """

    workbook = load_workbook(
        filename=str(file_path),
        read_only=True,
        data_only=True,
    )

    docs: list[Document] = []

    try:
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            rows_text: list[str] = []

            for row in sheet.iter_rows(values_only=True):
                values: list[str] = []

                for value in row:
                    if value is None:
                        values.append("")
                    else:
                        values.append(str(value).strip())

                if any(cell for cell in values):
                    rows_text.append(" | ".join(values))

            content = "\n".join(rows_text).strip()

            if not content:
                continue

            docs.append(
                Document(
                    page_content=f"Excel sheet: {sheet_name}\n\n{content}",
                    metadata={
                        "source": str(file_path),
                        "source_file": file_path.name,
                        "doc_type": _guess_doc_type(file_path.name),
                        "file_type": file_path.suffix.lower(),
                        "sheet_name": sheet_name,
                    },
                )
            )

    finally:
        workbook.close()

    return docs


def load_documents() -> list[Document]:
    """
    Load all supported documents from the documents folder.

    Currently the folder comes from:
    settings.pdf_dir

    Even if the name is pdf_dir, it can now contain:
    - PDF
    - TXT
    - Markdown
    - JSON
    - CSV
    - DOCX
    - Excel
    """

    documents_dir = _resolve_path(settings.pdf_dir)
    documents_dir.mkdir(parents=True, exist_ok=True)

    docs: list[Document] = []

    # 1. Load PDF files
    pdf_loader = PyPDFDirectoryLoader(str(documents_dir))
    pdf_docs = pdf_loader.load()

    for doc in pdf_docs:
        source = doc.metadata.get("source", "")

        doc.metadata["doc_type"] = _guess_doc_type(source)
        doc.metadata["source_file"] = Path(source).name if source else "unknown"
        doc.metadata["file_type"] = ".pdf"

    docs.extend(pdf_docs)

    # 2. Load other supported file types
    for file_path in documents_dir.rglob("*"):
        if not file_path.is_file():
            continue

        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            continue

        if suffix in {".txt", ".md"}:
            docs.extend(_load_text_file(file_path))

        elif suffix == ".json":
            docs.extend(_load_json_file(file_path))

        elif suffix == ".csv":
            docs.extend(_load_csv_file(file_path))

        elif suffix == ".docx":
            docs.extend(_load_docx_file(file_path))

        elif suffix in {".xlsx", ".xlsm", ".xltx", ".xltm"}:
            docs.extend(_load_excel_file(file_path))

    return docs


def split_documents(docs: list[Document]) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )

    return splitter.split_documents(docs)


def reset_vector_store() -> None:
    global _VECTOR_STORE

    _VECTOR_STORE = None
    chroma_dir = _resolve_path(settings.chroma_dir)

    if chroma_dir.exists():
        shutil.rmtree(chroma_dir)

    chroma_dir.mkdir(parents=True, exist_ok=True)


def ingest_pdfs(reset: bool = True) -> dict[str, Any]:
    """
    Ingest documents into ChromaDB.

    The function name stays ingest_pdfs() so main.py does not need changes.

    But now it supports:
    - PDF
    - TXT
    - Markdown
    - JSON
    - CSV
    - DOCX
    - Excel
    """

    if reset:
        reset_vector_store()

    docs = load_documents()
    chunks = split_documents(docs)

    if not chunks:
        return {
            "message": "No document content found. Add PDF, TXT, MD, JSON, CSV, DOCX, or Excel files and try again.",
            "loaded_documents": 0,
            "created_chunks": 0,
            "chroma_dir": str(_resolve_path(settings.chroma_dir)),
            "collection_name": settings.collection_name,
        }

    vector_store = get_vector_store()
    vector_store.add_documents(chunks)

    return {
        "message": "Documents ingested successfully.",
        "loaded_documents": len(docs),
        "created_chunks": len(chunks),
        "chroma_dir": str(_resolve_path(settings.chroma_dir)),
        "collection_name": settings.collection_name,
    }


def _format_context(docs: list[Document]) -> str:
    blocks: list[str] = []

    for i, doc in enumerate(docs, start=1):
        source_file = doc.metadata.get("source_file") or doc.metadata.get("source") or "unknown"
        page = doc.metadata.get("page")
        doc_type = doc.metadata.get("doc_type", "document")
        file_type = doc.metadata.get("file_type", "unknown")

        blocks.append(
            f"[Source {i}] file={source_file}, page={page}, type={doc_type}, file_type={file_type}\n"
            f"{doc.page_content}"
        )

    return "\n\n".join(blocks)


def _build_sources(docs: list[Document]) -> list[dict[str, Any]]:
    sources = []

    for doc in docs:
        sources.append(
            {
                "source": doc.metadata.get("source_file") or doc.metadata.get("source"),
                "page": doc.metadata.get("page"),
                "doc_type": doc.metadata.get("doc_type"),
                "preview": doc.page_content[:350].replace("\n", " ").strip(),
                "metadata": doc.metadata,
            }
        )

    return sources


def retrieve_documents(
    question: str,
    k: int | None = None,
    filters: dict[str, Any] | None = None,
) -> list[Document]:
    retrieval_started = time.perf_counter()
    vector_store = get_vector_store()

    search_kwargs: dict[str, Any] = {
        "k": k or settings.top_k,
    }

    if filters:
        search_kwargs["filter"] = filters

    retriever = vector_store.as_retriever(search_kwargs=search_kwargs)
    docs = retriever.invoke(question)
    print(f"[timing] retrieval time: {time.perf_counter() - retrieval_started:.3f}s")

    return docs


def _chunk_text(chunk: Any) -> str:
    content = getattr(chunk, "content", "")

    if isinstance(content, str):
        return content

    if content:
        return str(content)

    return ""


def ask(
    question: str,
    k: int | None = None,
    filters: dict[str, Any] | None = None,
    voice_mode: bool = False,
) -> dict[str, Any]:
    docs = retrieve_documents(question=question, k=k, filters=filters)

    if not docs:
        return {
            "answer": "Profood does not have enough data yet. Try ingesting documents first with POST /ingest.",
            "sources": [],
        }

    context = _format_context(docs)
    prompt = VOICE_PROMPT if voice_mode else PROMPT

    chain = prompt | get_llm()
    generation_started = time.perf_counter()

    response = chain.invoke(
        {
            "context": context,
            "question": question,
        }
    )
    print(f"[timing] total LLM generation time: {time.perf_counter() - generation_started:.3f}s")

    return {
        "answer": response.content,
        "sources": _build_sources(docs),
    }


def stream_answer_chunks(
    question: str,
    k: int | None = None,
    filters: dict[str, Any] | None = None,
    voice_mode: bool = False,
    stop_event: Event | None = None,
) -> Iterator[dict[str, Any]]:
    stream_started = time.perf_counter()
    docs = retrieve_documents(question=question, k=k, filters=filters)

    if not docs:
        yield {
            "type": "chunk",
            "text": "Profood does not have enough data yet. Try ingesting documents first with POST /ingest.",
        }
        yield {"type": "sources", "sources": []}
        return

    context = _format_context(docs)
    prompt = VOICE_PROMPT if voice_mode else PROMPT
    chain = prompt | get_llm()
    generation_started = time.perf_counter()
    first_token_logged = False

    for chunk in chain.stream(
        {
            "context": context,
            "question": question,
        }
    ):
        if stop_event and stop_event.is_set():
            break

        text = _chunk_text(chunk)

        if not text:
            continue

        if not first_token_logged:
            first_token_logged = True
            print(f"[timing] first token time: {time.perf_counter() - stream_started:.3f}s")

        yield {
            "type": "chunk",
            "text": text,
        }

    print(f"[timing] total LLM generation time: {time.perf_counter() - generation_started:.3f}s")

    if not stop_event or not stop_event.is_set():
        yield {
            "type": "sources",
            "sources": _build_sources(docs),
        }
